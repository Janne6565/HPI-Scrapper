# Author: Janne Keipert
# TODO: Login HPI
# TODO: Gather Sourcecode

from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.actions.mouse_button import MouseButton
import os
from datetime import date, timedelta
import docx
import progressbar
import Informations
from time import sleep


today = date.today()

options = webdriver.ChromeOptions()
options.headless = True

browser = webdriver.Chrome(options=options)
browser.get("https://www.hpi-schul-cloud.de/login")
listLinks = []

def login(): # Login into the Cloud
    print("Logging in")
    loginName = browser.find_element_by_xpath('//*[@id="name"]')
    loginKey = browser.find_element_by_xpath('//*[@id="password"]')
    submit = browser.find_element_by_xpath('//*[@id="submit-login"]')

    loginName.send_keys(Informations.hpiEmail)
    loginKey.send_keys(Informations.hpiPassword)
    submit.submit()
    if browser.current_url == "https://hpi-schul-cloud.de/dashboard":
        return True
    else:
        return False

def scrap(): # Gather the needed Information
    print("Starting Scrapping")
    listExercisesObj = browser.find_element_by_xpath('//*[@id="main-content"]/section[2]/div[2]')
    listExercisesList = listExercisesObj.text.split('\n')
    listExercisesOrdered = []
    listObjectsLinks = browser.find_elements_by_class_name("homework-link")
    for i in listObjectsLinks:
        listLinks.append(i.get_attribute("href"))
    for i in range(int((len(listExercisesList)-1) / 3)):
        listExercisesOrdered.append([listExercisesList[i * 3 + 2], listExercisesList[i * 3], listExercisesList[i * 3 + 1]])
    return listExercisesOrdered

def writeFile(listExercises): # Write the Informations in the File
    count = 0
    os.system('shutdown /s /t 999 /c "Du hast ' + str(len(listExercises)) + ' neue Aufgaben"')
    sleep(10)
    os.system('shutdown /a')
    with progressbar.ProgressBar(max_value=len(listExercises), enable_colors=False) as bar:
        listText = ["\n"]
        for i in listExercises:
            date = i[2].split()
            date = date[1]
            if date == "Abgabedatum" or date == "einem":
                pass
            else:
                subject = i[1].replace(" ", "").replace("10.4", "").replace("-", "").replace("Cordes", "").replace("fürden10.Jahrgang" ,"")
                exercise = i[0].replace(" ", "_").replace("10.4", "").replace("-", "").replace(",", "")
                dateExercise = today + timedelta(days=int(date))
                dateExercise = str(dateExercise).replace(" ", "").replace(".", "")
                fileWriteExercise = "Aufgaben/" + subject + "/" + dateExercise + "_" + exercise + ".docx"
                os.makedirs(os.path.dirname(fileWriteExercise), exist_ok=True)
                browser.get(listLinks[count])
                elementInstpect = browser.find_element_by_xpath('//*[@id="extended"]/div[1]/div[1]/div')

                textExercise = elementInstpect.text
                doc = docx.Document()
                doc.add_paragraph(textExercise)
                doc.add_paragraph(listLinks[count])
                listText.append("Aufgabe: " + exercise + " Fach: " + subject + " Abgabe: " + dateExercise)
                for ii in browser.find_elements_by_xpath('//*[@id="extended"]/div[1]/div[2]/ul/li/a'):
                    doc.add_paragraph(ii.get_attribute("href"))
                    print(ii.get_attribute("href"))
                if not (os.path.isfile(fileWriteExercise)):
                    doc.save(fileWriteExercise)
            bar.update(count)
            count += 1
        for i in listText:
            print(i)

def writeToDo(listExercises):
    browser.get("https://todoist.com/app")
    while True:
        try:
            browser.find_element_by_xpath('//*[@id="email"]').send_keys(Informations.todoistEmail)
            browser.find_element_by_xpath('//*[@id="password"]').send_keys(Informations.todoistPassword + "\n")
            break
        except:
            pass
    print("Worked")
    sleep(3)
    sleep(1)
    browser.find_element_by_xpath('//*[@id="projects_list"]/li/table/tbody/tr/td[3]').click()
    sleep(0.1)
    browser.find_element_by_xpath('//*[@id="menu_delete_text"]').click()
    sleep(0.1)
    browser.find_element_by_class_name('ist_button_red').click()
    sleep(0.1)
    browser.find_element_by_xpath('//*[@id="list_holder"]/div[2]/header/div/button').click()
    sleep(0.1)
    browser.find_element_by_xpath('//*[@id="edit_project_modal_field_name"]').send_keys("Schule")
    browser.find_element_by_class_name('ist_button_red').click()
    count = 0
    for i in listExercises:
        obj = browser.find_elements_by_xpath('//*[@id="labels_list"]/li/a')
        exist = False
        lesson = i[1].replace(" ", "_").replace("ü", "ue").replace("ä", "ae").replace("ö","oe").replace("Ü", "ue").replace("Ä", "Ae").replace("Ö", "oe")
        for ii in obj:
            print(ii.get_attribute("href"))
            if ii.get_attribute("href").__contains__(lesson):
                exist = True

        if not exist:
            browser.find_element_by_xpath('//*[@id="list_holder"]/div[3]/header/div/button').click()
            sleep(0.3)
            browser.find_element_by_xpath('//*[@id="edit_label_modal_field_name"]').send_keys(lesson + "\n")
        browser.find_element_by_xpath('//*[@id="projects_list"]/li[1]').click()
        sleep(1)
        browser.find_element_by_xpath('//*[@id="editor"]/div[2]/div/div/ul/li/div/div/ul/li/button').click()
        sleep(0.3)
        browser.find_element_by_xpath('//*[@id="editor"]/div[2]/div/div/ul/li/div/div/ul/li/form/div[1]/div[1]/div/div/div/div/div/div').send_keys("@" + lesson + " " + i[0])
        sleep(0.2)
        if (i[2] != "Kein Abgabedatum festgelegt") and not i[2].__contains__("Monat"):
            browser.find_element_by_xpath('//*[@id="editor"]/div[2]/div/div/ul/li/div/div/ul/li/form/div[1]/div[2]/div[1]/button[1]').click()
            sleep(0.3)
            browser.find_element_by_class_name("scheduler-input").find_element_by_tag_name("input").send_keys(i[2].replace("-", ".") + "\n")
        browser.find_element_by_class_name('quick_note_action').click()
        browser.find_element_by_xpath('//*[@id="quick_comment_input"]').send_keys(listLinks[count])
        browser.find_element_by_class_name('close_btn').click()
        browser.find_element_by_xpath('//*[@id="editor"]/div[2]/div/div/ul/li/div/div/ul/li/form/div[2]/button[1]').click()
        sleep(1)
        count += 1;


if login():
    print("Succesfull Login")
    listExercices = scrap()
    print(listExercices)
    if Informations.toDo:
        writeToDo(listExercices)
    writeFile(listExercices)
    input("Succesfully Finished press Enter to quit")
    browser.quit()
else:
    print("Du hast eine falsche EMail adresse oder Passwort angegeben")
    browser.quit()
